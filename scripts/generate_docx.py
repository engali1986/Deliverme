from docx import Document
from docx.shared import Pt
import os

root = r"C:\Users\Engal\Desktop\Deliverme\frontend"
output_path = r"C:\Users\Engal\Desktop\Deliverme\Deliverme_frontend_documentation.docx"

print('Walking', root)

doc = Document()
doc.add_heading('Deliverme Frontend - Files Tree and Code', level=1)

# Files tree
doc.add_heading('Files Tree', level=2)
exclude_dirs = {'node_modules', '.expo', 'android', 'ios', '__pycache__'}
for dirpath, dirnames, filenames in os.walk(root):
    # Skip excluded directories in-place to avoid walking them
    dirnames[:] = [d for d in dirnames if d not in exclude_dirs]
    rel = os.path.relpath(dirpath, root)
    if rel == '.':
        display = os.path.basename(root.rstrip(os.sep)) or 'frontend'
        indent = 0
    else:
        display = os.path.basename(dirpath)
        indent = rel.count(os.sep) + 1
    doc.add_paragraph('  ' * indent + display + '/')
    for fn in sorted(filenames):
        doc.add_paragraph('  ' * (indent + 1) + fn)

# Add a page break
from docx.enum.text import WD_BREAK
p = doc.add_paragraph()
p.add_run().add_break(WD_BREAK.PAGE)

# Add file contents
doc.add_heading('Files and Code', level=2)
all_files = []
for dirpath, dirnames, filenames in os.walk(root):
    for fn in filenames:
        fp = os.path.join(dirpath, fn)
        all_files.append(fp)

allowed_exts = {'.js', '.json', '.md', '.txt', '.mjs', '.jsx', '.ts', '.tsx', '.html', '.css', '.env', '.py', '.jsonc'}
all_files = sorted(all_files)
for fp in all_files:
    rel = os.path.relpath(fp, root)
    _, ext = os.path.splitext(fp)
    if ext.lower() not in allowed_exts:
        # Skip binary or large asset files
        doc.add_heading(rel + ' (skipped - binary/asset)', level=3)
        continue
    doc.add_heading(rel, level=3)
    try:
        with open(fp, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
    except Exception as e:
        content = f'<<Could not read file: {e}>>'
    # Sanitize content: remove NULL bytes and most control chars
    content = content.replace('\x00', '')
    filtered = ''.join(ch if (ch == '\n' or ch == '\t' or ord(ch) >= 32) else ' ' for ch in content)
    # Add code block in manageable chunks to avoid XML limitations
    para = doc.add_paragraph()
    max_chunk = 1000
    for i in range(0, len(filtered), max_chunk):
        chunk = filtered[i:i+max_chunk]
        run = para.add_run(chunk)
        font = run.font
        font.name = 'Courier New'
        font.size = Pt(9)
    doc.add_paragraph()

# Save document
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print('Saved:', output_path)
